"""
天气系统分析报告生成模块（第二份 Word）
======================================
生成包含天气系统描述文字和图片的分析报告。
"""
import os
import logging
from datetime import datetime
from typing import List, Dict

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from config.settings import REPORT_DIR, PLOT_LEVELS

logger = logging.getLogger(__name__)


class AnalysisReportGenerator:
    """天气系统分析报告生成器"""

    def __init__(self, current_time: datetime, init_hour: int):
        self.current_time = current_time
        self.init_hour = init_hour
        os.makedirs(REPORT_DIR, exist_ok=True)

    def generate(
        self,
        trough_analyses: List[Dict],
        vortex_analyses: List[Dict] = None,
        cold_vortex_analyses: List[Dict] = None,
        subtropical_high_analyses: List[Dict] = None,
        analysis_figures: Dict = None,
        low_level_jet_analysis: Dict = None,
    ) -> str:
        """
        生成分析报告。

        Args:
            trough_analyses:  槽线分析结果列表
            analysis_figures: 分析图片路径
                {level: {"obs": path, "fcst_tracking": path}, ...}

        Returns:
            报告文件路径
        """
        doc = Document()

        # 标题
        title = (
            f"天气系统分析报告 — "
            f"{self.current_time.strftime('%Y年%m月%d日')}"
            f"{self.init_hour:02d}时起报"
        )
        doc.add_heading(title, level=0)

        # 分析时间
        p = doc.add_paragraph()
        run = p.add_run(
            f"分析时间：{self.current_time.strftime('%Y年%m月%d日%H时%M分')}"
        )
        run.font.size = Pt(10)
        run.bold = True
        self._set_font(run, "微软雅黑")

        doc.add_paragraph("")

        # ── 一、影响天津的天气系统列表 ──
        doc.add_heading("一、未来12小时影响天津的天气系统", level=1)
        affecting_list = self._build_affecting_list(trough_analyses, vortex_analyses, cold_vortex_analyses, subtropical_high_analyses, low_level_jet_analysis)
        self._add_styled_paragraph(doc, affecting_list)

        # ── 二、系统描述 ──
        doc.add_heading("二、系统描述", level=1)

        doc.add_heading("2.1 实况描述", level=2)
        obs_text = self._build_obs_description(trough_analyses, vortex_analyses, cold_vortex_analyses, subtropical_high_analyses, low_level_jet_analysis)
        self._add_styled_paragraph(doc, obs_text)

        doc.add_heading("2.2 预报描述", level=2)
        fcst_text = self._build_fcst_description(trough_analyses, vortex_analyses, cold_vortex_analyses, subtropical_high_analyses, low_level_jet_analysis)
        self._add_styled_paragraph(doc, fcst_text)

        # ── 三、可视化（按层次） ──
        doc.add_heading("三、可视化", level=1)

        if analysis_figures:
            for level in PLOT_LEVELS:
                level_figs = analysis_figures.get(level, {})
                if not level_figs:
                    continue

                doc.add_heading(f"{level}hPa", level=2)

                obs_path = level_figs.get("obs", "")
                fcst_path = level_figs.get("fcst_tracking", "")

                self._add_image_pair(doc, obs_path, fcst_path)
                doc.add_paragraph("")

        # 保存
        filename = (
            f"天气系统分析报告_"
            f"{self.current_time.strftime('%Y%m%d')}_"
            f"{self.init_hour:02d}时起报.docx"
        )
        filepath = os.path.join(REPORT_DIR, filename)
        doc.save(filepath)
        logger.info(f"分析报告已保存: {filepath}")
        return filepath

    def _build_affecting_list(self, trough_analyses: List[Dict],
                               vortex_analyses: List[Dict] = None,
                               cold_vortex_analyses: List[Dict] = None,
                               subtropical_high_analyses: List[Dict] = None,
                               low_level_jet_analysis: Dict = None) -> str:
        systems = []

        for a in trough_analyses:
            if a["is_new"]:
                systems.append(f"新生槽{a['trough_num']}")
            else:
                systems.append(f"高空槽{a['trough_num']}")

        if vortex_analyses:
            for a in vortex_analyses:
                if a["is_new"]:
                    systems.append(f"新生低涡{a['vortex_num']}（{a['level']}hPa）")
                else:
                    systems.append(f"低涡{a['vortex_num']}（{a['level']}hPa）")

        if cold_vortex_analyses:
            for a in cold_vortex_analyses:
                if a["is_new"]:
                    systems.append(f"新生{a['vortex_type']}{a['cv_num']}")
                else:
                    systems.append(f"{a['vortex_type']}{a['cv_num']}")

        if subtropical_high_analyses:
            for a in subtropical_high_analyses:
                if a.get("is_affecting"):
                    systems.append("副热带高压")
                    break

        if low_level_jet_analysis and low_level_jet_analysis.get("is_affecting"):
            systems.append("低空急流（850hPa）")

        if not systems:
            return "未来12小时内无天气系统影响天津。"

        return "、".join(systems)

    def _build_obs_description(self, trough_analyses: List[Dict],
                                vortex_analyses: List[Dict] = None,
                                cold_vortex_analyses: List[Dict] = None,
                                subtropical_high_analyses: List[Dict] = None,
                                low_level_jet_analysis: Dict = None) -> str:
        lines = []

        # 高空槽
        existing_troughs = [a for a in trough_analyses if not a["is_new"]]
        for a in existing_troughs:
            lines.append(
                f"高空槽{a['trough_num']}（500hPa）："
                f"目前位于{a.get('location', '未知位置')}，"
                f"是{a['tilt_type']}槽"
            )

        # 低涡
        if vortex_analyses:
            existing_vortex = [a for a in vortex_analyses if not a["is_new"]]
            for a in existing_vortex:
                lines.append(
                    f"低涡{a['vortex_num']}（{a['level']}hPa）："
                    f"目前低压中心位于{a.get('location', '未知位置')}，"
                    f"低涡中心强度为{a.get('strength', '未知')}dagpm"
                )

        # 冷涡
        if cold_vortex_analyses:
            existing_cv = [a for a in cold_vortex_analyses if not a["is_new"]]
            for a in existing_cv:
                lines.append(
                    f"{a['vortex_type']}{a['cv_num']}（500hPa）："
                    f"中心强度为{a.get('strength', '未知')}dagpm"
                )

        # 副高
        if subtropical_high_analyses:
            obs_items = [a for a in subtropical_high_analyses if a["type"] == "obs" and a.get("is_affecting")]
            for a in obs_items:
                lines.append(f"副热带高压（500hPa）：{a['description']}")

        # 低空急流（第一个影响时次为实况时）
        if (low_level_jet_analysis and low_level_jet_analysis.get("is_affecting")
                and low_level_jet_analysis.get("first_source") == "obs"):
            lines.append(f"低空急流（850hPa）：{low_level_jet_analysis['wording']}")

        if not lines:
            return "当前实况中无天气系统影响天津。"

        return "\n".join(lines)
    
    def _build_fcst_description(self, trough_analyses: List[Dict],
                                 vortex_analyses: List[Dict] = None,
                                 cold_vortex_analyses: List[Dict] = None,
                                 subtropical_high_analyses: List[Dict] = None,
                                 low_level_jet_analysis: Dict = None) -> str:
        lines = []
        num = 1

        # 高空槽
        for a in trough_analyses:
            prefix = "新生槽" if a["is_new"] else "高空槽"
            lines.append(
                f"{num}. {prefix}{a['trough_num']}（500hPa）："
                f"未来将自{a['from_dir']}向{a['to_dir']}移动，"
                f"预计{a['impact_time']}影响天津"
            )
            num += 1

        # 低涡
        if vortex_analyses:
            for a in vortex_analyses:
                prefix = "新生低涡" if a["is_new"] else "低涡"
                lines.append(
                    f"{num}. {prefix}{a['vortex_num']}（{a['level']}hPa）："
                    f"未来将自{a['from_dir']}向{a['to_dir']}移动，"
                    f"天津位于低涡{a.get('quadrant', '未知')}象限，"
                    f"预计{a['impact_time']}影响天津"
                )
                num += 1

        # 冷涡
        if cold_vortex_analyses:
            for a in cold_vortex_analyses:
                prefix = f"新生{a['vortex_type']}" if a["is_new"] else a["vortex_type"]
                lines.append(
                    f"{num}. {prefix}{a['cv_num']}（500hPa）："
                    f"未来将自{a['from_dir']}向{a['to_dir']}移动，"
                    f"天津位于冷涡{a.get('quadrant', '未知')}象限，"
                    f"预计{a['impact_time']}影响天津"
                )
                num += 1

        # 副高
        if subtropical_high_analyses:
            fcst_items = [a for a in subtropical_high_analyses if a["type"] == "fcst" and a.get("is_affecting")]
            for a in fcst_items:
                lines.append(f"{num}. 副热带高压（500hPa）：{a['description']}")
                num += 1

        # 低空急流（第一个影响时次为预报时）
        if (low_level_jet_analysis and low_level_jet_analysis.get("is_affecting")
                and low_level_jet_analysis.get("first_source") == "fcst"):
            lines.append(
                f"{num}. 低空急流（850hPa）："
                f"预计{low_level_jet_analysis['impact_time']}"
                f"{low_level_jet_analysis['wording']}"
            )
            num += 1

        if not lines:
            return "未来12小时内无天气系统影响天津。"

        return "\n".join(lines)
    
    def _add_styled_paragraph(self, doc, text: str):
        """添加带字体样式的段落"""
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(10)
        self._set_font(run, "微软雅黑")

    def _add_image_pair(self, doc: Document, obs_path: str, fcst_path: str):
        """在文档中添加一对并排图片（通过表格实现）"""
        table = doc.add_table(rows=2, cols=2)
        table.autofit = True

        # 第一行：标题
        titles = ["实况分析", "预报追踪"]
        for col in range(2):
            p = table.cell(0, col).paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(titles[col])
            run.font.size = Pt(10)
            run.bold = True
            self._set_font(run, "微软雅黑")

        # 第二行：图片
        img_width = Inches(3.0)

        if obs_path and os.path.exists(obs_path):
            p = table.cell(1, 0).paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(obs_path, width=img_width)
        else:
            table.cell(1, 0).text = "（图片缺失）"

        if fcst_path and os.path.exists(fcst_path):
            p = table.cell(1, 1).paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(fcst_path, width=img_width)
        else:
            table.cell(1, 1).text = "（图片缺失）"

    @staticmethod
    def _set_font(run, font_name: str):
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)