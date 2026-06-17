"""
Word 报告生成模块
=================
将所有可视化图片按层级结构组装成 Word 文档。

文档结构:
    标题: 天气系统识别报告
    一、实况
        {时间标签}
            {level}hPa 层次:
                ┌──────────────┬──────────────┐
                │ 要素场图片    │ 系统识别图片  │
                └──────────────┴──────────────┘
    二、预报
        {时间标签} (+000h)
            500hPa / 700hPa / 850hPa ...
        {时间标签} (+003h)
            ...

每组两张图片通过表格实现并排显示。

依赖:
    python-docx (pip install python-docx)
"""
import os
import logging
from datetime import datetime
from typing import Dict, Tuple
from collections import defaultdict

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from config.settings import REPORT_DIR, REPORT_CONFIG, PLOT_LEVELS

logger = logging.getLogger(__name__)


class WordReportGenerator:
    """
    Word 报告生成器

    将 figure_paths 字典中的图片按照规定的层级结构组装成 .docx 文件。
    """

    def __init__(self, current_time: datetime, init_hour: int):
        self.current_time = current_time
        self.init_hour = init_hour
        os.makedirs(REPORT_DIR, exist_ok=True)

    def generate(
        self, figure_paths: Dict[Tuple[str, str, int], Tuple[str, str, str]]
    ) -> str:
        """
        生成 Word 报告。

        Args:
            figure_paths: 图片路径字典
                key:   (data_type, time_label, level)
                value: (field_fig_path, system_fig_path, system_names)

        Returns:
            生成的 Word 文件路径
        """
        doc = Document()

        # 标题
        title = (
            f"天气系统识别报告 — "
            f"{self.current_time.strftime('%Y年%m月%d日')}"
            f"{self.init_hour:02d}时起报"
        )
        doc.add_heading(title, level=0)

        # 按 data_type 分组
        obs_items = {}
        fcst_items = defaultdict(dict)

        for (dtype, time_label, level), paths in figure_paths.items():
            if dtype == "obs":
                obs_items[level] = paths
            else:
                fcst_items[time_label][level] = paths

        # ── 一、实况 ──
        doc.add_heading("一、实况", level=1)
        if obs_items:
            # 取实况时间标签
            obs_time_label = None
            for (dtype, tl, _) in figure_paths:
                if dtype == "obs":
                    obs_time_label = tl
                    break
            if obs_time_label:
                doc.add_heading(obs_time_label, level=2)

            for level in PLOT_LEVELS:
                if level not in obs_items:
                    continue
                field_path, system_path, system_names = obs_items[level]
                doc.add_heading(f"{level}hPa 层次", level=3)
                self._add_image_pair(doc, field_path, system_path, system_names)
        else:
            doc.add_paragraph("无实况数据")

        # ── 二、预报 ──
        doc.add_heading("二、预报", level=1)
        if fcst_items:
            # 按时间标签排序
            sorted_times = sorted(fcst_items.keys())
            for time_label in sorted_times:
                doc.add_heading(time_label, level=2)
                levels_data = fcst_items[time_label]

                for level in PLOT_LEVELS:
                    if level not in levels_data:
                        continue
                    field_path, system_path, system_names = levels_data[level]
                    doc.add_heading(f"{level}hPa 层次", level=3)
                    self._add_image_pair(doc, field_path, system_path, system_names)
        else:
            doc.add_paragraph("无预报数据")

        # 保存
        filename = (
            f"天气系统识别报告_"
            f"{self.current_time.strftime('%Y%m%d')}_"
            f"{self.init_hour:02d}时起报.docx"
        )
        filepath = os.path.join(REPORT_DIR, filename)
        doc.save(filepath)
        logger.info(f"Word 报告已保存: {filepath}")
        return filepath

    def _add_image_pair(
        self, doc: Document, field_path: str, system_path: str, system_names: str = "无"
    ):
        """
        在文档中添加一对并排图片（通过表格实现）。

        Args:
            doc:         Document 对象
            field_path:  要素场图片路径
            system_path: 系统识别图片路径
        """
        table = doc.add_table(rows=2, cols=2)
        table.autofit = True

        # 第一行: 标题
        from docx.oxml.ns import qn

        titles = ["风场+高度场", f"系统识别（{system_names}）"]
        for col in range(2):
            p = table.cell(0, col).paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(titles[col])
            run.font.size = Pt(9)
            run.bold = True
            run.font.name = "微软雅黑"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

        # 第二行: 图片
        img_width = Inches(REPORT_CONFIG.image_width_inches)

        if field_path and os.path.exists(field_path):
            p = table.cell(1, 0).paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(field_path, width=img_width)
        else:
            table.cell(1, 0).text = "（图片缺失）"

        if system_path and os.path.exists(system_path):
            p = table.cell(1, 1).paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(system_path, width=img_width)
        else:
            table.cell(1, 1).text = "（图片缺失）"

        # 表格后加空行
        doc.add_paragraph("")
